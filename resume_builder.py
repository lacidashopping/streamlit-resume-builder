import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO

def create_resume_docx(data):
    """Generates a structured .docx file with a profile photo and dynamic sections."""
    doc = Document()
    
    # Page setup adjustments (1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Header with Photo Layout (Using a hidden table for side-by-side structure)
    if data['photo']:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)
        
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        p_name = left_cell.paragraphs[0]
        r_name = p_name.add_run(data['name'])
        r_name.font.size = Pt(24)
        r_name.bold = True
        
        p_contact = left_cell.add_paragraph()
        p_contact.add_run(f"{data['email']}  |  {data['phone']}\n{data['location']}\n")
        if data['linkedin']:
            p_contact.add_run(f"LinkedIn: {data['linkedin']}\n")
        if data['github']:
            p_contact.add_run(f"GitHub: {data['github']}")
            
        p_photo = right_cell.paragraphs[0]
        p_photo.alignment = 2
        p_photo.add_run().add_picture(data['photo'], width=Inches(1.5))
        
        doc.add_paragraph()
    else:
        title = doc.add_paragraph()
        title.alignment = 1
        run_name = title.add_run(data['name'])
        run_name.font.size = Pt(24)
        run_name.bold = True
        
        contact_p = doc.add_paragraph()
        contact_p.alignment = 1
        contact_p.add_run(f"{data['email']}  |  {data['phone']}  |  {data['location']}\n")
        if data['linkedin']:
            contact_p.add_run(f"LinkedIn: {data['linkedin']}  ")
        if data['github']:
            contact_p.add_run(f"|  GitHub: {data['github']}")

    def add_section_heading(text):
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
        doc.add_paragraph("―" * 50).paragraph_format.space_after = Pt(6)

    # 2. Professional Summary
    if data['summary']:
        add_section_heading("Professional Summary")
        doc.add_paragraph(data['summary'])

    # 3. Dynamic Work Experience
    if data['experience']:
        add_section_heading("Work Experience")
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_role = p.add_run(f"{exp['role']} ")
            r_role.bold = True
            p.add_run(f"at {exp['company']} ({exp['duration']})")
            
            if exp['description']:
                for bullet in exp['description'].split("\n"):
                    if bullet.strip():
                        doc.add_paragraph(bullet.strip(), style='List Bullet')

    # 4. Education
    if data['education']:
        add_section_heading("Education")
        for edu in data['education']:
            p = doc.add_paragraph()
            r_deg = p.add_run(f"{edu['degree']} ")
            r_deg.bold = True
            p.add_run(f"― {edu['school']} ({edu['year']})")

    # 5. Skills
    if data['skills']:
        add_section_heading("Skills")
        doc.add_paragraph(data['skills'])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- Streamlit Session State Initialization ---
if "exp_count" not in st.session_state:
    st.session_state.exp_count = 1

st.set_page_config(page_title="Dynamic Resume Builder", layout="centered")

# --- 🚀 RAMERLABS PROMO BANNER ---
st.info("💡 **Looking for a custom application like this one?** We build custom web tools, dashboards, and automated systems tailored to your business needs. Visit us at [ramerlabs.com](https://ramerlabs.com) to launch your project today!")

st.title("📄 Dynamic Resume Builder")
st.subheader("Add multiple roles and upload your 2x2 photo layout")

# --- 🚀 SIDEBAR AD (Stays visible as they work) ---
with st.sidebar:
    st.markdown("### 🛠️ Need a Custom App?")
    st.markdown(
        """
        Do you need a bespoke internal tool, dynamic document generator, database system, or client portal?
        
        **We build software for businesses.**
        
        👉 **[Visit ramerlabs.com](https://ramerlabs.com)**
        """
    )
    st.divider()

# 1. Personal & Media Details Block
st.header("1. Personal Details")
uploaded_photo = st.file_uploader("Upload 2x2 ID Photo (JPG/PNG)", type=["jpg", "png", "jpeg"])

col1, col2 = st.columns(2)
with col1:
    name = st.text_input("Full Name", "John Doe")
    email = st.text_input("Email Address", "johndoe@example.com")
with col2:
    phone = st.text_input("Phone Number", "+1 (555) 019-2834")
    location = st.text_input("Location", "New York, NY")
    
linkedin = st.text_input("LinkedIn URL")
github = st.text_input("GitHub URL")

# 2. Professional Summary
st.header("2. Professional Summary")
summary = st.text_area("Write your summary intro here", "Highly motivated professional...")

# 3. Dynamic Work Experience Form Generation
st.header("3. Work Experience")

experiences = []
for i in range(st.session_state.exp_count):
    st.markdown(f"#### Job Position #{i+1}")
    ec1, ec2, ec3 = st.columns([2, 2, 2])
    with ec1:
        role = st.text_input(f"Job Title", key=f"role_{i}")
    with ec2:
        company = st.text_input(f"Company", key=f"company_{i}")
    with ec3:
        duration = st.text_input(f"Duration (e.g., 2023 - Present)", key=f"dur_{i}")
        
    desc = st.text_area(f"Job Description (One bullet per line)", key=f"desc_{i}")
    
    if role or company:
        experiences.append({"role": role, "company": company, "duration": duration, "description": desc})
    st.markdown("---")

cb1, cb2 = st.columns([1, 4])
with cb1:
    if st.button("➕ Add Job"):
        st.session_state.exp_count += 1
        st.rerun()
with cb2:
    if st.button("➖ Remove Last Job") and st.session_state.exp_count > 1:
        st.session_state.exp_count -= 1
        st.rerun()

# 4. Education & Skills Blocks
st.header("4. Education")
edu_degree = st.text_input("Degree/Certification", "B.S. in Computer Science")
edu_school = st.text_input("School/University", "State University")
edu_year = st.text_input("Graduation Year", "2022")

st.header("5. Core Skills")
skills = st.text_area("Skills (Comma-separated)", "Python, SQL, Project Management")

# Compile Execution Block
st.markdown("### Generate Resume")
if st.button("🚀 Compile and Prepare Document"):
    
    resume_data = {
        "name": name, "email": email, "phone": phone, "location": location,
        "linkedin": linkedin, "github": github, "summary": summary, "skills": skills,
        "photo": uploaded_photo,
        "experience": experiences,
        "education": [{"school": edu_school, "degree": edu_degree, "year": edu_year}]
    }
    
    docx_buffer = create_resume_docx(resume_data)
    
    st.success("🎉 Document generated ready below!")
    st.download_button(
        label="📥 Download Resume as .docx",
        data=docx_buffer,
        file_name=f"{name.replace(' ', '_')}_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
